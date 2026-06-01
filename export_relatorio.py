# -*- coding: utf-8 -*-  # define a codificação do ficheiro para suportar acentos

from datetime import date, timedelta  # importa date e timedelta para trabalhar com datas
from collections import defaultdict  # importa defaultdict para dicionários com valor padrão
import math  # importa math para cálculos numéricos e arredondamentos
import re  # importa re para expressões regulares (tratamento de texto de roles etc.)
import calendar

try:  # tenta importar as classes do python-docx
    from docx import Document  # importa Document para criar ficheiros .docx
    from docx.shared import Pt, RGBColor  # importa Pt e RGBColor para tamanhos e cores de fonte
    DOCX_DISPONIVEL = True  # se o import funcionar, marcamos que o DOCX está disponível
except ImportError:  # se der erro ao importar python-docx
    DOCX_DISPONIVEL = False  # marcamos que o DOCX não está disponível

def clean_role(role: str) -> str:  # define função utilitária para "normalizar" o nome da função (role)
    role = role.strip()  # remove espaços em branco no início e no fim da string
    role = role.lower()  # converte o texto para minúsculas para comparar de forma consistente
    return role  # devolve a string normalizada

# =========================================================
# ========  PARTE 5 — RELATÓRIOS / ANÁLISE GERAL  =========
# =========================================================

def contar_horas_turno(turno):
    """Recebe string 'HH:MM–HH:MM' e devolve horas (float)."""
    ini, fim = turno.split("–")
    h1, m1 = map(int, ini.split(":"))
    h2, m2 = map(int, fim.split(":"))
    dt1 = timedelta(hours=h1, minutes=m1)
    dt2 = timedelta(hours=h2, minutes=m2)
    diff = (dt2 - dt1).total_seconds() / 3600
    if diff < 0:
        diff += 24
    return diff

def analisar_necessidades_e_equipa(estado, escala_4_semanas):

    """
    Faz análise geral:
    - Necessidades totais de cada role
    - Quantidade de funcionários por role
    - Escassez / excesso real (considerando que 1 funcionário pode trabalhar vários dias)
    - Horas efetivamente trabalhadas por funcionário
    - Recomendações gerais
    Retorna um dicionário completo.
    """

    e = estado
    roles = e["roles"]
    employees = e["employees"]

    # Regras da empresa/legislação (limite semanal base)
    regras = e.get("rules", {})
    max_horas_semana_default = float(regras.get("max_horas_semana", 40.0))

    # Número de semanas efetivamente analisadas (4, 5, ...)
    semanas_periodo = len(escala_4_semanas) or 1

    # -------------------------------------------------------
    # 1) Contabilizar horas necessárias por role na escala gerada
    # -------------------------------------------------------
    horas_necessarias = {r: 0.0 for r in roles}

    for semana in escala_4_semanas:
        for d, turnos_dict in semana.items():
            for turno, itens in turnos_dict.items():
                h_turno = contar_horas_turno(turno)
                for it in itens:
                    if isinstance(it, tuple):
                        _, role = it
                        # garante que a role existe no dicionário
                        if role not in horas_necessarias:
                            horas_necessarias[role] = 0.0
                        horas_necessarias[role] += h_turno
                    else:
                        # VAGA ABERTA também conta como necessidade
                        m = re.search(r"\((.*?)\)", str(it))
                        if m:
                            role = m.group(1).strip()
                            if role not in horas_necessarias:
                                horas_necessarias[role] = 0.0
                            horas_necessarias[role] += h_turno

    # -------------------------------------------------------
    # 2) Contar funcionários disponíveis por role
    # -------------------------------------------------------
    func_por_role = {r: [] for r in roles}

    for emp_id, emp in employees.items():
        for r in emp.get("roles", []):
            if r in func_por_role:
                func_por_role[r].append(emp_id)

    # -------------------------------------------------------
    # 3) Calcular horas trabalhadas por funcionário
    # -------------------------------------------------------
    horas_trabalhadas = {emp_id: 0.0 for emp_id in employees}

    for semana in escala_4_semanas:
        for d, turnos_dict in semana.items():
            for turno, itens in turnos_dict.items():
                h_turno = contar_horas_turno(turno)
                for it in itens:
                    if isinstance(it, tuple):
                        emp_id, _ = it
                        if emp_id in horas_trabalhadas:
                            horas_trabalhadas[emp_id] += h_turno

    # -------------------------------------------------------
    # 4) Identificar Escassez e Excesso REAL de staff por role
    # -------------------------------------------------------
    analise_roles = {}

    for r in roles:
        qtde_func = len(func_por_role.get(r, []))
        horas_role = float(horas_necessarias.get(r, 0.0))

        # capacidade total do staff naquela role, com base nas regras reais
        capacidade_teorica = 0.0
        for emp_id in func_por_role.get(r, []):
            emp = employees.get(emp_id, {})
            horas_sem_emp = float(emp.get("max_horas_semana_emp", max_horas_semana_default))
            capacidade_teorica += horas_sem_emp * semanas_periodo

        if horas_role > capacidade_teorica:
            deficit = horas_role - capacidade_teorica
            status = "DEFICIT"
        else:
            deficit = capacidade_teorica - horas_role
            status = "EXCESSO"

        analise_roles[r] = {
            "role": r,
            "staff": qtde_func,
            "horas_necessarias": horas_role,
            "capacidade_teorica": capacidade_teorica,
            "status": status,
            "margem": deficit,
        }

    # -------------------------------------------------------
    # 5) Identificar recomendações neste relatório
    # -------------------------------------------------------
    recomendacoes = []

    for r, info in analise_roles.items():
        if info["status"] == "DEFICIT":
            recomendacoes.append(
                f"⚠️ A função **{r.title()}** apresenta falta de staff. "
                f"Necessita aprox. {info['margem']:.1f} horas adicionais no período."
            )
        else:
            recomendacoes.append(
                f"ℹ️ A função **{r.title()}** está com margem de segurança "
                f"de aproximadamente {info['margem']:.1f} horas no período."
            )

    # Recomendar equilíbrio
    recomendacoes.append(
        "\n💡 **Sugestão Geral:** redistribuir polivalentes para roles deficitárias "
        "ou contratar reforço caso a distribuição atual continue gerando vagas abertas."
    )

    return {
        "horas_necessarias": horas_necessarias,
        "func_por_role": func_por_role,
        "horas_trabalhadas": horas_trabalhadas,
        "analise_roles": analise_roles,
        "recomendacoes": recomendacoes,
    }

def analisar_4_semanas(estado, escala_4_semanas, primeira_segunda, ano_alvo, mes_alvo):
    """
    Analisa o PERÍODO da escala (agora alinhado a um mês), usando todas
    as semanas geradas em `escala_4_semanas` (pode ser 4, 5, ...).
    Devolve (lista_de_semanas, metrics) para o relatório DOCX.
    """
    employees = estado.get("employees", {})  # obtém dicionário de colaboradores
    regras = estado.get("rules", {})  # obtém regras trabalhistas
    max_horas_semana = regras.get("max_horas_semana", 40.0)  # limite de horas semanais

    # número de semanas efetivamente geradas (4, 5, ...)
    semanas_periodo = len(escala_4_semanas)
    if semanas_periodo <= 0:
        semanas_periodo = 1  # evita divisões por zero

    # datas de início e fim do PERÍODO lógico (mês escolhido)
    _, last_day_num = calendar.monthrange(ano_alvo, mes_alvo)
    inicio_geral = date(ano_alvo, mes_alvo, 1)
    fim_geral = date(ano_alvo, mes_alvo, last_day_num)

    role_stats = defaultdict(
        lambda: {
            "horas_necessarias": 0.0,
            "horas_alocadas": 0.0,
            "horas_vagas": 0.0,
            "vagas_abertas_slots": 0,
        }
    )

    emp_stats = {}
    for emp_id, emp in employees.items():
        emp_stats[emp_id] = {
            "nome": emp.get("nome", emp_id),
            "roles": [clean_role(r) for r in emp.get("roles", [])],
            "horas_total": 0.0,
            "horas_matinal": 0.0,
            "horas_diurno": 0.0,
            "horas_noturno": 0.0,
            "dias_trabalhados": set(),
            "turnos_total": 0,
        }

    def classificar_periodo(turno_str):
        try:
            ini_str, _ = turno_str.split("–")
        except ValueError:
            return "diurno"
        try:
            hora_ini = int(ini_str.split(":")[0])
        except ValueError:
            return "diurno"

        if 5 <= hora_ini < 12:
            return "matinal"
        elif 12 <= hora_ini < 20:
            return "diurno"
        else:
            return "noturno"

    # percorre TODAS as semanas geradas
    for semana in escala_4_semanas:
        for dia, turnos_dict in semana.items():
            for turno_str, itens in turnos_dict.items():
                horas_turno = contar_horas_turno(turno_str)
                periodo = classificar_periodo(turno_str)

                for it in itens:
                    if isinstance(it, tuple):
                        emp_id, role_raw = it
                        role_norm = clean_role(role_raw)
                        rs = role_stats[role_norm]

                        rs["horas_necessarias"] += horas_turno
                        rs["horas_alocadas"] += horas_turno

                        emp_info = emp_stats.get(emp_id)
                        if emp_info is not None:
                            emp_info["horas_total"] += horas_turno
                            emp_info["turnos_total"] += 1
                            emp_info["dias_trabalhados"].add(dia)

                            if periodo == "matinal":
                                emp_info["horas_matinal"] += horas_turno
                            elif periodo == "diurno":
                                emp_info["horas_diurno"] += horas_turno
                            else:
                                emp_info["horas_noturno"] += horas_turno
                    else:
                        texto = str(it)
                        m = re.search(r"\((.*?)\)", texto)
                        if m:
                            role_raw = m.group(1).strip()
                            role_norm = clean_role(role_raw)
                        else:
                            role_norm = clean_role("vaga_aberta")

                        rs = role_stats[role_norm]
                        rs["horas_necessarias"] += horas_turno
                        rs["horas_vagas"] += horas_turno
                        rs["vagas_abertas_slots"] += 1

    # capacidade teórica por função agora ajustada ao nº de semanas
    capacidade_teorica = defaultdict(float)

    for emp_id, emp in employees.items():
        horas_sem = float(emp.get("max_horas_semana_emp", max_horas_semana))
        horas_periodo = horas_sem * semanas_periodo

        for r in emp.get("roles", []):
            role_norm = clean_role(r)
            capacidade_teorica[role_norm] += horas_periodo

    metrics = {
        "role_stats": {role: stats for role, stats in role_stats.items()},
        "emp_stats": emp_stats,
        "capacidade_teorica": dict(capacidade_teorica),
        "primeira_segunda": primeira_segunda,
        "num_semanas": semanas_periodo,
        "periodo_inicio": inicio_geral,
        "periodo_fim": fim_geral,
    }

    return escala_4_semanas, metrics

def _add_paragrafo(doc, texto, bold=False):  # função auxiliar para adicionar parágrafos ao documento Word
    p = doc.add_paragraph(texto)             # cria um novo parágrafo no documento com o texto indicado
    if bold:                                 # se o parâmetro "bold" for True
        for run in p.runs:                   # percorre todos os "runs" (pedaços de texto) do parágrafo
            run.bold = True                  # marca esse pedaço de texto como negrito
    return p                                 # devolve o parágrafo criado (caso seja útil para mais formatações depois)

# =========================
# RELATÓRIO EM WORD (.DOCX)
# =========================

def gerar_relatorio_docx(estado, semanas, metrics):
    """Gera um ficheiro .docx com análise da escala do PERÍODO (mês)."""

    if not DOCX_DISPONIVEL:
        print("⚠️ python-docx não está instalado. Relatório .docx não será gerado.")
        return

    role_stats = metrics["role_stats"]
    emp_stats = metrics["emp_stats"]
    capacidade_teorica = metrics["capacidade_teorica"]
    primeira_segunda = metrics["primeira_segunda"]

    regras = estado["rules"]
    max_horas_semana = regras.get("max_horas_semana", 40.0)

    semanas_periodo = metrics.get("num_semanas", 4)
    inicio_geral = metrics.get("periodo_inicio", primeira_segunda)
    fim_geral = metrics.get(
        "periodo_fim",
        primeira_segunda + timedelta(days=semanas_periodo * 7 - 1),
    )

    colaboradores_por_role = defaultdict(int)
    for emp_id, emp in estado.get("employees", {}).items():
        for r in emp.get("roles", []):
            role_norm = clean_role(r)
            colaboradores_por_role[role_norm] += 1

    doc = Document()

    # título principal
    h1 = doc.add_heading("RELATÓRIO OPERACIONAL – ESCALA DO PERÍODO", level=1)
    h1.alignment = 0

    _add_paragrafo(
        doc,
        f"Período analisado: {inicio_geral.strftime('%d/%m/%Y')} a {fim_geral.strftime('%d/%m/%Y')}",
    )
    _add_paragrafo(doc, f"Número de semanas no período: {semanas_periodo}")
    _add_paragrafo(doc, "Legislação: Portugal – Código do Trabalho")
    _add_paragrafo(
        doc,
        f"Limite de horas semanais por colaborador: {max_horas_semana:.1f} h",
    )
    _add_paragrafo(doc, "")

    # ---------- 1. Visão por função ----------
    doc.add_heading("1. Visão por função (role)", level=2)
    _add_paragrafo(
        doc,
        "Nesta secção são comparadas as horas necessárias, horas escaladas e vagas\n"
        "em aberto por função. Ajuda a identificar funções em falta ou com demasiada folga.",
    )
    _add_paragrafo(doc, "")

    if not role_stats:
        _add_paragrafo(doc, "Nenhuma função foi configurada.")
        _add_paragrafo(doc, "")
    else:
        for role_norm, dados in sorted(role_stats.items(), key=lambda x: x[0]):
            horas_nec_total = dados.get("horas_necessarias", 0.0)
            if horas_nec_total <= 0:
                continue

            horas_alo_total = dados.get("horas_alocadas", 0.0)
            horas_vag_total = dados.get("horas_vagas", 0.0)
            vagas_slots = dados.get("vagas_abertas_slots", 0)

            cap_total = capacidade_teorica.get(role_norm, 0.0)
            horas_nec_semana = horas_nec_total / semanas_periodo if semanas_periodo > 0 else 0.0

            if max_horas_semana > 0 and semanas_periodo > 0:
                pessoas_equivalentes_nec = horas_nec_total / (max_horas_semana * semanas_periodo)
                pessoas_equivalentes_disp = cap_total / (max_horas_semana * semanas_periodo) if cap_total > 0 else 0.0
            else:
                pessoas_equivalentes_nec = 0.0
                pessoas_equivalentes_disp = 0.0

            cap_semana = cap_total / semanas_periodo if cap_total > 0 and semanas_periodo > 0 else 0.0
            n_colabs = colaboradores_por_role.get(role_norm, 0)
            nome_legivel = role_norm.title()

            _add_paragrafo(doc, f"Função: {nome_legivel}", bold=True)
            _add_paragrafo(
                doc,
                f"  • Colaboradores configurados com esta função: {n_colabs}",
            )
            _add_paragrafo(
                doc,
                f"  • Horas necessárias no período: {horas_nec_total:.1f} h "
                f"(~{horas_nec_semana:.1f} h/semana)",
            )
            _add_paragrafo(
                doc,
                f"  • Horas efetivamente escaladas: {horas_alo_total:.1f} h",
            )

            if horas_vag_total > 0.01:
                _add_paragrafo(
                    doc,
                    f"  • Horas que ficaram em 'VAGA ABERTA': {horas_vag_total:.1f} h "
                    f"(slots em falta: {vagas_slots})",
                )
            else:
                _add_paragrafo(
                    doc,
                    "  • Não foram registadas vagas abertas nesta função.",
                )

            if cap_total > 0:
                _add_paragrafo(
                    doc,
                    f"  • Capacidade teórica para esta função no período: {cap_total:.1f} h "
                    f"(~{cap_semana:.1f} h/semana)",
                )
                _add_paragrafo(
                    doc,
                    f"  • Equivalente a aproximadamente {pessoas_equivalentes_nec:.2f} pessoas a tempo inteiro "
                    f"(necessidade) e {pessoas_equivalentes_disp:.2f} a tempo inteiro (disponíveis).",
                )

                if horas_vag_total > 0.01:
                    _add_paragrafo(
                        doc,
                        "  → Há falta de cobertura em alguns turnos desta função. "
                        "Pode ser necessário contratar ou reforçar estas horas.",
                    )
                else:
                    if cap_total > 0 and horas_nec_total < 0.6 * cap_total:
                        _add_paragrafo(
                            doc,
                            "  → A equipa desta função aparenta estar folgada em relação às necessidades médias.",
                        )
                    else:
                        _add_paragrafo(
                            doc,
                            "  → A carga desta função está razoavelmente alinhada com a capacidade da equipa.",
                        )
            else:
                _add_paragrafo(
                    doc,
                    "  • Não existem colaboradores configurados com esta função, apesar de haver horas necessárias.",
                )

            _add_paragrafo(doc, "")

    # ---------- 2. Visão por colaborador ----------
    doc.add_heading("2. Visão por colaborador", level=2)
    _add_paragrafo(
        doc,
        "Resumo de horas e turnos por colaborador, ajudando a ver quem está mais\n"
        "carregado e quem está com menos utilização.",
    )
    _add_paragrafo(doc, "")

    if not emp_stats:
        _add_paragrafo(doc, "Nenhum colaborador foi configurado.")
        _add_paragrafo(doc, "")
    else:
        for emp_id, dados in sorted(emp_stats.items(), key=lambda x: x[1]["nome"]):
            nome = dados.get("nome", emp_id)
            roles_emp = ", ".join(dados.get("roles", [])) or "—"
            horas_total = dados.get("horas_total", 0.0)
            horas_matinal = dados.get("horas_matinal", 0.0)
            horas_diurno = dados.get("horas_diurno", 0.0)
            horas_noturno = dados.get("horas_noturno", 0.0)
            dias_trabalhados = len(dados.get("dias_trabalhados", set()))
            turnos_total = dados.get("turnos_total", 0)

            horas_semanais_med = horas_total / semanas_periodo if semanas_periodo > 0 else 0.0
            horas_por_dia = horas_total / dias_trabalhados if dias_trabalhados > 0 else 0.0

            _add_paragrafo(doc, f"Colaborador: {nome}", bold=True)
            _add_paragrafo(doc, f"  • Funções registadas: {roles_emp}")
            _add_paragrafo(
                doc,
                f"  • Horas totais no período: {horas_total:.1f} h "
                f"(~{horas_semanais_med:.1f} h/semana; {dias_trabalhados} dias diferentes)",
            )
            _add_paragrafo(
                doc,
                f"  • Horas por tipo de período – Matinal: {horas_matinal:.1f} h, "
                f"Diurno: {horas_diurno:.1f} h, Noturno: {horas_noturno:.1f} h",
            )
            _add_paragrafo(
                doc,
                f"  • Turnos atribuídos: {turnos_total} (média ~{horas_por_dia:.1f} h por dia trabalhado)",
            )

            if horas_semanais_med > max_horas_semana * 0.95:
                _add_paragrafo(
                    doc,
                    "  → Atenção: este colaborador está muito próximo do limite semanal definido.",
                )
            elif horas_semanais_med < max_horas_semana * 0.4:
                _add_paragrafo(
                    doc,
                    "  → Nota: este colaborador está a trabalhar bem abaixo do limite semanal.",
                )

            _add_paragrafo(doc, "")

    # ---------- 3. Pontos críticos ----------
    doc.add_heading("3. Pontos críticos e recomendações", level=2)
    _add_paragrafo(
        doc,
        "Resumo das funções em falta ou com sobra de capacidade, com sugestões\n"
        "simples de melhoria.",
    )
    _add_paragrafo(doc, "")

    funcoes_em_falta = []
    funcoes_com_sobra = []

    for role_norm, dados in role_stats.items():
        horas_nec = dados.get("horas_necessarias", 0.0)
        if horas_nec <= 0:
            continue
        horas_alo = dados.get("horas_alocadas", 0.0)
        horas_vag = dados.get("horas_vagas", 0.0)
        cap_total = capacidade_teorica.get(role_norm, 0.0)

        if horas_vag > 0.01:
            funcoes_em_falta.append((role_norm, horas_vag, horas_nec, horas_alo))
        if cap_total > 0 and horas_nec < 0.6 * cap_total:
            sobra = cap_total - horas_nec
            funcoes_com_sobra.append((role_norm, sobra, horas_nec, horas_alo))

    funcoes_em_falta.sort(key=lambda x: x[1], reverse=True)
    funcoes_com_sobra.sort(key=lambda x: x[1], reverse=True)

    if not funcoes_em_falta and not funcoes_com_sobra:
        _add_paragrafo(
            doc,
            "Não foram detetados desequilíbrios muito fortes entre necessidade e capacidade.\n"
            "Ainda assim recomenda-se uma revisão manual da escala, sobretudo em dias críticos.",
        )
    else:
        if funcoes_em_falta:
            _add_paragrafo(doc, "3.1 Funções com falta de capacidade", bold=True)
            for role_norm, horas_vag, horas_nec, horas_alo in funcoes_em_falta:
                nome_legivel = role_norm.title()
                _add_paragrafo(
                    doc,
                    f"  • {nome_legivel}: {horas_vag:.1f} h em falta "
                    f"(necessárias: {horas_nec:.1f} h, escaladas: {horas_alo:.1f} h).",
                )
            _add_paragrafo(
                doc,
                "   Recomendações: avaliar contratação, reforço de horas ou redistribuição\n"
                "   realista de funções, garantindo que não se viola descanso nem limites legais.",
            )
            _add_paragrafo(doc, "")

        if funcoes_com_sobra:
            _add_paragrafo(doc, "3.2 Funções com capacidade sobrante", bold=True)
            for role_norm, horas_sobra, horas_nec, horas_alo in funcoes_com_sobra:
                nome_legivel = role_norm.title()
                _add_paragrafo(
                    doc,
                    f"  • {nome_legivel}: ~{horas_sobra:.1f} h de capacidade sobrante "
                    f"(necessárias: {horas_nec:.1f} h, escaladas: {horas_alo:.1f} h).",
                )
            _add_paragrafo(
                doc,
                "   Recomendações: ponderar redução de horas nesta função ou deslocar parte\n"
                "   desta capacidade para funções em falta.",
            )

    _add_paragrafo(doc, "")
    _add_paragrafo(
        doc,
        "Este relatório foi gerado automaticamente com base na escala atual. Deve ser\n"
        "usado como apoio ao planeamento, sempre complementado com a leitura da realidade\n"
        "no terreno.",
    )

    nome_docx = f"Relatorio_Escala_{primeira_segunda.strftime('%Y%m%d')}.docx"
    try:
        doc.save(nome_docx)
        print(f"✅ Relatório DOCX gerado: {nome_docx}")
    except Exception as exc:
        print(f"⚠️ Erro ao gravar relatório DOCX: {exc}")
